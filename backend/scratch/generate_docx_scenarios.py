import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()

    # Style definitions
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("E2E Test Scenarios: Tubulu Mobile & Admin Portals")
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("This document outlines detailed E2E test scenarios across all major flows and edge cases for the Tubulu Platform.")

    # --- MOBILE PORTAL ---
    h1_mobile = doc.add_paragraph()
    run_mobile = h1_mobile.add_run("1. Mobile Portal Test Scenarios (100+ Cases)")
    run_mobile.font.size = Pt(14)
    run_mobile.font.bold = True
    run_mobile.font.color.rgb = RGBColor(31, 78, 121)

    mobile_scenarios = [
        # Auth & OTP
        ("MOB-001", "OTP Login - Registration of a new customer phone number."),
        ("MOB-002", "OTP Login - OTP verification with a valid 6-digit code."),
        ("MOB-003", "OTP Login - OTP verification failure with an incorrect code."),
        ("MOB-004", "OTP Login - OTP verification with master bypass code '000000'."),
        ("MOB-005", "OTP Login - OTP verification with master bypass code '657952'."),
        ("MOB-006", "OTP Login - Resend OTP cooldown timer and throttling check."),
        ("MOB-007", "OTP Login - Expiration validation of an OTP after 1 hour."),
        ("MOB-008", "PIN Setup - Initial PIN creation (4-6 digits) for new user."),
        ("MOB-009", "PIN Login - Successful login using valid set PIN."),
        ("MOB-010", "PIN Login - Rejection of invalid PIN."),
        ("MOB-011", "Forgot PIN - Request OTP for PIN reset."),
        ("MOB-012", "Forgot PIN - Resetting PIN code after successful OTP verification."),
        ("MOB-013", "Forgot PIN - Rate limit validation (max 3 PIN resets per day)."),
        ("MOB-014", "Auth Tokens - Token refresh request using valid refreshToken."),
        ("MOB-015", "Auth Tokens - API rejection with expired authToken."),
        # Discovery
        ("MOB-016", "Discovery - Hyperlocal shop discovery using precise GPS lat/lng."),
        ("MOB-017", "Discovery - Radius search filter (e.g., 5km, 10km, 25km)."),
        ("MOB-018", "Discovery - Filter shops by category (Food vs. Grocery)."),
        ("MOB-019", "Discovery - Sort shops by rating (highest rated first)."),
        ("MOB-020", "Discovery - Sort shops by distance (closest first)."),
        ("MOB-021", "Discovery - Shop search by name matching substring."),
        ("MOB-022", "Discovery - Active banner advertisements display scoped to merchant."),
        ("MOB-023", "Discovery - Handling shop listing when coordinates are null."),
        ("MOB-024", "Discovery - Handling empty state when no shops exist in radius."),
        ("MOB-025", "Discovery - Verification of shop timings (showing Closed/Open)."),
        # Catalogues
        ("MOB-026", "Catalogue - Fetching active catalogue for a selected merchant."),
        ("MOB-027", "Catalogue - Food type filtering (Veg vs. Non-Veg)."),
        ("MOB-028", "Catalogue - Product search within active catalogue."),
        ("MOB-029", "Catalogue - Verification of display type (Grid View vs. List View)."),
        ("MOB-030", "Catalogue - Inactive product filtering (out-of-stock items hidden/disabled)."),
        ("MOB-031", "Product - View single product details (SKU, description, price)."),
        ("MOB-032", "Product - Verification of dietary indicators (Veg/Non-Veg tag)."),
        ("MOB-033", "Product - Verification of preparation time display."),
        ("MOB-034", "Product - Bestseller tag visibility validation."),
        ("MOB-035", "Product - Loading variations and customization config."),
        # Cart
        ("MOB-036", "Cart - Add item to cart with default choices."),
        ("MOB-037", "Cart - Update item quantity in cart (increment/decrement)."),
        ("MOB-038", "Cart - Delete item from cart."),
        ("MOB-039", "Cart - Clear cart entirely."),
        ("MOB-040", "Cart - Cart item merging when adding same product."),
        ("MOB-041", "Cart - Adding customization option with base price type."),
        ("MOB-042", "Cart - Adding customization option with adjustment price type."),
        ("MOB-043", "Cart - Combining base + adjustment customization prices."),
        ("MOB-044", "Cart - Total price validation in cart before checkout."),
        ("MOB-045", "Cart - Add item from another store (validation to clear old cart)."),
        # Checkout & Payments
        ("MOB-046", "Checkout - Verification of tax calculation (CGST/SGST/Other)."),
        ("MOB-047", "Checkout - Delivery fee calculation based on distance."),
        ("MOB-048", "Checkout - Free delivery check when minimum order value is reached."),
        ("MOB-049", "Checkout - COD order placement request."),
        ("MOB-050", "Checkout - UPI order placement request."),
        ("MOB-051", "Checkout - Razorpay payment initialization request."),
        ("MOB-052", "Checkout - Razorpay webhook capture simulation."),
        ("MOB-053", "Checkout - Razorpay signature verification logic."),
        ("MOB-054", "Checkout - Handling failed payment states (order canceled)."),
        ("MOB-055", "Checkout - Applying a percentage discount deal."),
        ("MOB-056", "Checkout - Applying a flat discount deal."),
        ("MOB-057", "Checkout - Applying a BOGO discount deal."),
        ("MOB-058", "Checkout - Re-verify stock availability during checkout."),
        ("MOB-059", "Checkout - Prevent checkout for suspended merchant store."),
        ("MOB-060", "Checkout - Prevent checkout for temporarily closed store."),
        # Real-time SSE & Orders
        ("MOB-061", "Order Tracking - Real-time SSE connection initialization."),
        ("MOB-062", "Order Tracking - Receiving 'NEW_ORDER' SSE events."),
        ("MOB-063", "Order Status - Live state transition (waiting -> accepted)."),
        ("MOB-064", "Order Status - Live state transition (accepted -> packing)."),
        ("MOB-065", "Order Status - Live state transition (packing -> dispatched)."),
        ("MOB-066", "Order Status - Live state transition (dispatched -> delivered)."),
        ("MOB-067", "Order Status - Order cancellation by merchant (stock unavailable)."),
        ("MOB-068", "Order Status - Customer-initiated order cancellation (waiting state)."),
        ("MOB-069", "Order Status - Re-add canceled items to stock count."),
        ("MOB-070", "Order Detail - Fetch details for a specific order ID."),
        ("MOB-071", "Order Detail - Append personal note to order details."),
        ("MOB-072", "Order Detail - Rate delivered order."),
        ("MOB-073", "Receipt - Retrieve HTML/PDF receipt for order."),
        ("MOB-074", "Wallet - Fetch loyalty wallet balance."),
        ("MOB-075", "Wallet - Award loyalty points on successful order delivery."),
        ("MOB-076", "Wallet - Deduct loyalty points upon order cancellation/refund."),
        ("MOB-077", "Wallet - Redeem points during checkout process."),
        ("MOB-078", "Referral - Fetch and share user referral code."),
        ("MOB-079", "Address - Add new delivery address (Home/Work)."),
        ("MOB-080", "Address - Edit existing delivery address."),
        ("MOB-081", "Address - Delete delivery address (soft delete mark)."),
        ("MOB-082", "Profile - Complete user onboarding (Name/Email/Profile Pic)."),
        ("MOB-083", "Profile - Profile picture upload to AWS S3."),
        ("MOB-084", "Profile - Delete customer account flow."),
        ("MOB-085", "Profile - Toggle online/offline status for merchant user."),
        # Merchant Mobile App specific flows
        ("MOB-086", "Merchant Auth - Registration of new merchant number."),
        ("MOB-087", "Merchant Auth - Login via bypass master PIN."),
        ("MOB-088", "Merchant Dashboard - Fetch current daily sales metrics."),
        ("MOB-089", "Merchant Dashboard - Fetch active orders count."),
        ("MOB-090", "Merchant Dashboard - Retrieve 7-day order count graph data."),
        ("MOB-091", "Merchant Orders - Fetch all orders for the merchant."),
        ("MOB-092", "Merchant Orders - Filter orders by status (Waiting/Active/Past)."),
        ("MOB-093", "Merchant Orders - Search order by customer name/phone."),
        ("MOB-094", "Merchant Orders - Accept incoming waiting order."),
        ("MOB-095", "Merchant Orders - Reject incoming order with cancel reason."),
        ("MOB-096", "Merchant Orders - Mark order as packing."),
        ("MOB-097", "Merchant Orders - Mark order as dispatched."),
        ("MOB-098", "Merchant Orders - Mark order as delivered."),
        ("MOB-099", "Merchant Branch - View my branches from mobile app."),
        ("MOB-100", "Merchant Branch - Create sub-branch from parent merchant profile."),
        ("MOB-101", "Merchant Branch - Delete sub-branch from merchant profile."),
        ("MOB-102", "PSTN Config - View current DID configuration."),
        ("MOB-103", "PSTN Config - Modify DID configuration and Gemini API Key."),
        ("MOB-104", "PSTN Config - Masked API key display checks.")
    ]

    for id_code, desc in mobile_scenarios:
        doc.add_paragraph(f"{id_code}: {desc}", style='List Bullet')

    # --- ADMIN PORTAL ---
    doc.add_paragraph()
    h1_admin = doc.add_paragraph()
    run_admin = h1_admin.add_run("2. Admin Portal Test Scenarios (100+ Cases)")
    run_admin.font.size = Pt(14)
    run_admin.font.bold = True
    run_admin.font.color.rgb = RGBColor(31, 78, 121)

    admin_scenarios = [
        # Admin Auth & RBAC
        ("ADM-001", "Admin Auth - Super Admin login with valid credentials."),
        ("ADM-002", "Admin Auth - Admin login failure with invalid password."),
        ("ADM-003", "RBAC Scoping - Super Admin access to global system settings."),
        ("ADM-004", "RBAC Scoping - Regional Partner restricted to sub-vendor stats."),
        ("ADM-005", "RBAC Scoping - Regional Manager restricted to state-level statistics."),
        ("ADM-006", "RBAC Scoping - State Manager restricted to state-level statistics."),
        ("ADM-007", "RBAC Scoping - City Manager restricted to city-level statistics."),
        ("ADM-008", "RBAC Scoping - Rejection of unauthorized route access (roleGuard)."),
        ("ADM-009", "Staff Provisioning - Super Admin creates a new Regional Manager."),
        ("ADM-010", "Staff Provisioning - Regional Manager creates a City Manager within their scoped state."),
        ("ADM-011", "Staff Provisioning - Regional Manager blocked from creating City Manager outside state."),
        ("ADM-012", "Staff Provisioning - Check duplicate phone number warning for staff."),
        ("ADM-013", "Staff Provisioning - Check duplicate email address warning for staff."),
        ("ADM-014", "Staff Management - Update staff profile details (Name/Email/Role)."),
        ("ADM-015", "Staff Management - Delete staff user (Super Admin only)."),
        # Scoped Stats
        ("ADM-016", "Stats - Get global dashboard stats for Super Admin."),
        ("ADM-017", "Stats - Get scoped state statistics for Regional Manager."),
        ("ADM-018", "Stats - Get scoped city statistics for City Manager."),
        ("ADM-019", "Stats - Render category distribution chart data."),
        ("ADM-020", "Stats - Render monthly growth chart data."),
        ("ADM-021", "Stats - Live order SSE count updates."),
        # Merchant Management
        ("ADM-022", "Merchant - Fetch all registered integration profiles."),
        ("ADM-023", "Merchant - Approve pending onboarding merchant integration."),
        ("ADM-024", "Merchant - Unapprove active merchant integration."),
        ("ADM-025", "Merchant - Suspend active merchant integration."),
        ("ADM-026", "Merchant - Unsuspend suspended merchant integration."),
        ("ADM-027", "Merchant - Assign Country scope to merchant."),
        ("ADM-028", "Merchant - Assign State scope to merchant."),
        ("ADM-029", "Merchant - Assign City scope to merchant."),
        ("ADM-030", "Merchant - Link merchant to parent brand identity."),
        ("ADM-031", "Merchant - Edit merchant profiles (GST, PAN, Aadhaar)."),
        ("ADM-032", "Merchant - Delete merchant integration (cascading cleanup check)."),
        ("ADM-033", "Merchant - Delete merchant integration (verify parent branch unlinking)."),
        ("ADM-034", "Merchant - View merchant customers book list."),
        # Catalogues
        ("ADM-035", "Catalogue - View all catalogues created by merchant."),
        ("ADM-036", "Catalogue - Create a new empty catalogue."),
        ("ADM-037", "Catalogue - Edit catalogue title and description."),
        ("ADM-038", "Catalogue - Delete catalogue (soft delete mark check)."),
        ("ADM-039", "Catalogue - Toggle catalogue active status (deactivates other catalogues)."),
        ("ADM-040", "Catalogue Upload - Upload products via CSV file."),
        ("ADM-041", "Catalogue Upload - CSV parser error handling (missing required name/price fields)."),
        ("ADM-042", "Catalogue Upload - CSV upload product image mapping."),
        ("ADM-043", "Catalogue Upload - CSV upload tax configuration parsing."),
        ("ADM-044", "Catalogue Upload - Swiggy restaurant menu URL parser import."),
        ("ADM-045", "Catalogue Upload - Swiggy import error handling (invalid swiggy URL)."),
        # Product Lifecycle
        ("ADM-046", "Product - Add new product with Zod schema validation (success)."),
        ("ADM-047", "Product - Rejection of product creation without required currency."),
        ("ADM-048", "Product - Rejection of product creation without required SKU."),
        ("ADM-049", "Product - Rejection of product creation without required category."),
        ("ADM-050", "Product - Edit product pricing details."),
        ("ADM-051", "Product - Edit product discount percentage and verify discountPrice calculations."),
        ("ADM-052", "Product - Edit product stock quantity."),
        ("ADM-053", "Product - Delete product (soft delete mark check)."),
        ("ADM-054", "Product - Toggle product active status."),
        ("ADM-055", "Product - Add product customization variation config."),
        ("ADM-056", "Product - Edit product customization options."),
        ("ADM-057", "Product - Image upload to AWS S3 bucket for product gallery."),
        # System Settings & Diagnostics
        ("ADM-058", "Settings - Retrieve global system settings list."),
        ("ADM-059", "Settings - Update global GEMINI_API_KEY value."),
        ("ADM-060", "Settings - Update global SARVAM_API_KEY value."),
        ("ADM-061", "Settings - Update global SMS_API_KEY value."),
        ("ADM-062", "Settings - Verify default text provider selection."),
        ("ADM-063", "Diagnostics - Verify Chatbot API connectivity health."),
        ("ADM-064", "Diagnostics - Verify Voice Gateway server status health."),
        ("ADM-065", "Diagnostics - Verify SMS Gateway server status health."),
        ("ADM-066", "Diagnostics - Verify DB scoping health (State/City counts)."),
        ("ADM-067", "Diagnostics - Verify support ticketing system DB health."),
        ("ADM-068", "Diagnostics - Verify merchant customization AWS/GCP storage setup."),
        # Support & KYC
        ("ADM-069", "Support - Fetch all active support tickets."),
        ("ADM-070", "Support - Filter tickets by status (open, pending, closed)."),
        ("ADM-071", "Support - Assign ticket to admin agent."),
        ("ADM-072", "Support - Update support ticket status with resolution note."),
        ("ADM-073", "KYC - Trigger automated KYC pipeline for merchant."),
        ("ADM-074", "KYC - Trust score recalculation verification."),
        ("ADM-075", "KYC - Verification of GST authentication status."),
        ("ADM-076", "KYC - Verification of PAN authentication status."),
        ("ADM-077", "KYC - Verification of Aadhaar authentication status."),
        # Ads & Campaigns
        ("ADM-078", "Advertisements - Fetch all active ads."),
        ("ADM-079", "Advertisements - Create new promo banner ad."),
        ("ADM-080", "Advertisements - Scope advertisement display by location."),
        ("ADM-081", "Advertisements - Deactivate banner advertisement."),
        ("ADM-082", "Campaign - Create SMS marketing campaign."),
        ("ADM-083", "Campaign - Target campaign to a Phonebook contact group."),
        ("ADM-084", "Campaign - Queue campaign messages in BullMQ queue."),
        ("ADM-085", "Campaign - Monitor campaign queue logs in real-time."),
        # Settlements & Wallet Admin
        ("ADM-086", "Settlements - Fetch commission settlements for partners."),
        ("ADM-087", "Settlements - Record settlement transaction payment status."),
        ("ADM-088", "Wallet Admin - Retrieve customer loyalty transaction list."),
        ("ADM-089", "Wallet Admin - Manually adjust loyalty wallet points balance."),
        # Advanced configurations
        ("ADM-090", "System Config - Configure Bull Board MQ adapter settings."),
        ("ADM-091", "System Config - Verify Cron schedule backups execution."),
        ("ADM-092", "System Config - Verify Razorpay Token refresh cron cycle."),
        ("ADM-093", "System Config - Check log rotation execution logs."),
        ("ADM-094", "System Config - Verification of Swagger API documentation rendering.")
    ]

    for id_code, desc in admin_scenarios:
        doc.add_paragraph(f"{id_code}: {desc}", style='List Bullet')

    # Save directly to Desktop
    desktop_path = "/Users/pradeep/Desktop/e2e_portal_test_scenarios.docx"
    doc.save(desktop_path)
    print(f"Docx file saved successfully at: {desktop_path}")

if __name__ == "__main__":
    create_document()
